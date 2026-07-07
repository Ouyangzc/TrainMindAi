"""DOCX parser."""

from docx import Document

from app.models.kb import DocumentPage


def _table_to_rows(table) -> list[list[str]]:  # noqa: ANN001
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    return rows


async def parse_docx(
    file_path: str, document_id: int, document_version_id: int
) -> list[DocumentPage]:
    """Split DOCX content by Heading 1 paragraphs."""
    doc = Document(file_path)
    pages: list[DocumentPage] = []
    current_title: str | None = None
    current_lines: list[str] = []
    page_number = 1

    def flush_page() -> None:
        nonlocal page_number, current_lines, current_title
        text = "\n".join(line for line in current_lines if line).strip()
        if not text and not current_title:
            return
        pages.append(
            DocumentPage(
                document_id=document_id,
                document_version_id=document_version_id,
                page_number=page_number,
                title=current_title or f"第 {page_number} 页",
                text=text,
                tables_json=[_table_to_rows(table) for table in doc.tables],
            )
        )
        page_number += 1
        current_lines = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name == "Heading 1":
            flush_page()
            current_title = text or f"第 {page_number} 页"
            current_lines = []
        elif text:
            current_lines.append(text)

    flush_page()
    return pages
